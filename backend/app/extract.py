"""Extract plain text from uploaded files for use as LLM context.

Supports PDF, Word (.docx), and a broad range of plain-text / code / data
formats. Returns the extracted text plus a flag indicating truncation.
"""
from __future__ import annotations

import io
import logging
from typing import Tuple

from .config import settings

logger = logging.getLogger("chatbot.extract")

# Extensions we treat as UTF-8 plain text (decoded directly).
_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".yaml", ".yml",
    ".xml", ".html", ".htm", ".log", ".ini", ".cfg", ".toml", ".rst",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".h", ".cpp", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".scala", ".sh",
    ".bash", ".sql", ".r", ".m", ".pl", ".lua", ".dart", ".vue", ".svelte",
}


class ExtractionError(Exception):
    """Raised when a file cannot be turned into usable text."""


# Image extensions handled via the vision path (not text extraction).
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


def is_image(filename: str) -> bool:
    return _ext(filename) in _IMAGE_EXTS


def prepare_image(data: bytes) -> tuple[str, str]:
    """Downscale/recompress an uploaded image and return (base64_str, mime).

    Groq caps a base64 image at ~4MB and 33 megapixels, so we shrink the
    longest side and re-encode as JPEG, stepping quality/size down until it
    fits comfortably. Raises ExtractionError on unreadable images.
    """
    try:
        from PIL import Image
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("Image support is not installed on the server.") from exc

    import base64
    import io

    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception as exc:
        raise ExtractionError("Could not read this image (it may be corrupt).") from exc

    # Flatten transparency onto white so we can save as JPEG.
    if img.mode in ("RGBA", "LA", "P"):
        img = img.convert("RGBA")
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[-1])
        img = bg
    else:
        img = img.convert("RGB")

    # Downscale the longest side.
    max_dim = settings.max_image_dim
    if max(img.size) > max_dim:
        ratio = max_dim / max(img.size)
        img = img.resize((int(img.width * ratio), int(img.height * ratio)))

    limit_bytes = int(settings.max_image_b64_mb * 1024 * 1024)
    quality = 85
    while True:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        raw = buf.getvalue()
        # base64 inflates size by ~4/3; check the encoded length.
        if len(raw) * 4 / 3 <= limit_bytes or quality <= 35:
            if len(raw) * 4 / 3 > limit_bytes:
                # Still too big at low quality — shrink dimensions and retry.
                img = img.resize((int(img.width * 0.8), int(img.height * 0.8)))
                quality = 85
                continue
            break
        quality -= 15

    b64 = base64.b64encode(raw).decode("ascii")
    return b64, "image/jpeg"


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("PDF support is not installed on the server.") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError("Could not read this PDF (it may be corrupt).") from exc
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            txt = ""
        if txt.strip():
            pages.append(f"[Page {i}]\n{txt.strip()}")
    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractionError(
            "No selectable text found in this PDF. It may be a scanned image; "
            "OCR is not supported."
        )
    return text


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # pragma: no cover
        raise ExtractionError("DOCX support is not installed on the server.") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError("Could not read this Word document.") from exc
    lines = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include simple table text too.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines).strip()
    if not text:
        raise ExtractionError("This Word document appears to be empty.")
    return text


def _from_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Could not decode this file as text.")


def extract_text(filename: str, data: bytes) -> Tuple[str, bool]:
    """Return (text, truncated) for the uploaded file.

    Raises ExtractionError for unsupported or unreadable files.
    """
    if not data:
        raise ExtractionError("The uploaded file is empty.")

    ext = _ext(filename)
    if ext == ".pdf":
        text = _from_pdf(data)
    elif ext == ".docx":
        text = _from_docx(data)
    elif ext == ".doc":
        raise ExtractionError(
            "Legacy .doc files are not supported. Please save as .docx or PDF."
        )
    elif ext in _TEXT_EXTS or ext == "":
        text = _from_text(data)
    else:
        raise ExtractionError(
            f"Unsupported file type '{ext or 'unknown'}'. Supported: PDF, DOCX, "
            "and common text/code files."
        )

    text = text.strip()
    truncated = False
    if len(text) > settings.max_doc_chars:
        text = text[: settings.max_doc_chars]
        truncated = True
    return text, truncated